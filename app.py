import streamlit as st
import anthropic
from docx import Document
import io

st.set_page_config(page_title="영어 학원 선생님 도우미", page_icon="📚")

st.title("📚 영어 학원 선생님 도우미")
st.write("수업 계획서, 학부모 알림장, 숙제를 AI가 자동으로 만들어드려요!")

with st.sidebar:
    st.header("⚙️ 설정")
    api_key = st.text_input("Anthropic API 키", type="password")

tab1, tab2, tab3 = st.tabs(["📝 수업 계획서", "📱 학부모 알림장", "📄 숙제 제작"])

# ── 탭 1: 수업 계획서 ──
with tab1:
    st.subheader("📋 수업 정보 입력")
    col1, col2 = st.columns(2)
    with col1:
        grade = st.selectbox("학년", ["1학년", "2학년", "3학년", "4학년", "5학년", "6학년"])
        duration = st.selectbox("수업 시간", ["40분", "50분", "60분", "90분"])
    with col2:
        level = st.selectbox("학생 수준", ["초급 (왕초보)", "중급 (기초 가능)", "고급 (대화 가능)"])
        students = st.number_input("학생 수", min_value=1, max_value=30, value=5)
    topic = st.text_input("수업 주제", placeholder="예: 과일 이름 배우기, 날씨 표현, 자기소개 등")
    goal = st.text_area("수업 목표", placeholder="예: 과일 이름 10개를 영어로 말할 수 있다.", height=80)

    if st.button("📝 수업 계획서 생성하기", type="primary"):
        if not api_key:
            st.warning("왼쪽 사이드바에 API 키를 입력해주세요!")
        elif not topic:
            st.warning("수업 주제를 입력해주세요!")
        else:
            with st.spinner("AI가 수업 계획서를 작성 중이에요... ✍️"):
                client = anthropic.Anthropic(api_key=api_key)
                prompt = f"""
당신은 초등학교 영어 전문 강사입니다.
아래 정보를 바탕으로 상세한 수업 계획서를 한국어로 작성해주세요.
- 대상: 초등학교 {grade}
- 수업 시간: {duration}
- 학생 수준: {level}
- 학생 수: {students}명
- 수업 주제: {topic}
- 수업 목표: {goal}
다음 형식으로 작성해주세요:
1. 수업 개요
2. 학습 목표 (3가지)
3. 준비물
4. 단계별 수업 진행 계획 (도입 → 전개 → 정리)
5. 추천 활동 및 게임
6. 숙제 제안
초등학생 눈높이에 맞고 흥미롭게 작성해주세요!
"""
                message = client.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=2000,
                    messages=[{"role": "user", "content": prompt}]
                )
                result = message.content[0].text
                st.success("✅ 수업 계획서가 완성됐어요!")
                st.divider()
                st.markdown(result)
                st.download_button(
                    label="📥 수업 계획서 다운로드 (txt)",
                    data=result,
                    file_name=f"수업계획서_{topic}.txt",
                    mime="text/plain"
                )

# ── 탭 2: 학부모 알림장 ──
with tab2:
    st.subheader("📱 학부모 알림장 정보 입력")
    col3, col4 = st.columns(2)
    with col3:
        student_name = st.text_input("학생 이름", placeholder="예: 김민준")
        class_date = st.date_input("수업 날짜")
    with col4:
        notice_topic = st.text_input("오늘 수업 주제", placeholder="예: 과일 이름 배우기")
        homework = st.text_input("숙제 내용", placeholder="예: 과일 단어 10개 외워오기")
    next_class = st.text_input("다음 수업 예고", placeholder="예: 다음 시간엔 색깔 표현을 배울 거예요!")
    special_note = st.text_area("특이사항 (선택)", placeholder="예: 오늘 발표를 아주 잘했어요!", height=80)

    if st.button("📱 알림장 생성하기", type="primary"):
        if not api_key:
            st.warning("왼쪽 사이드바에 API 키를 입력해주세요!")
        elif not student_name or not notice_topic:
            st.warning("학생 이름과 수업 주제를 입력해주세요!")
        else:
            with st.spinner("AI가 알림장을 작성 중이에요... ✍️"):
                client = anthropic.Anthropic(api_key=api_key)
                prompt = f"""
당신은 친근하고 따뜻한 영어 학원 선생님입니다.
아래 정보를 바탕으로 카카오톡으로 보낼 학부모 알림장을 작성해주세요.
- 학생 이름: {student_name}
- 수업 날짜: {class_date}
- 오늘 수업 주제: {notice_topic}
- 숙제: {homework}
- 다음 수업 예고: {next_class}
- 특이사항: {special_note if special_note else "없음"}
다음 조건을 지켜주세요:
1. 카카오톡 메시지처럼 짧고 친근한 말투로 작성
2. 이모지를 적절히 사용해서 읽기 편하게
3. 오늘 수업 내용 요약, 숙제 안내, 다음 수업 예고 순서로 작성
4. 전체 길이는 10줄 이내로 간결하게
5. 마지막엔 따뜻한 인사로 마무리
"""
                message = client.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=1000,
                    messages=[{"role": "user", "content": prompt}]
                )
                result = message.content[0].text
                st.success("✅ 알림장이 완성됐어요!")
                st.divider()
                st.markdown(result)
                st.download_button(
                    label="📥 알림장 복사용 다운로드 (txt)",
                    data=result,
                    file_name=f"알림장_{student_name}_{class_date}.txt",
                    mime="text/plain"
                )

# ── 탭 3: 숙제 제작 ──
with tab3:
    st.subheader("📄 숙제 정보 입력")
    col5, col6 = st.columns(2)
    with col5:
        hw_grade = st.selectbox("학년", ["1학년", "2학년", "3학년", "4학년", "5학년", "6학년"], key="hw_grade")
        hw_level = st.selectbox("학생 수준", ["초급 (왕초보)", "중급 (기초 가능)", "고급 (대화 가능)"], key="hw_level")
    with col6:
        hw_topic = st.text_input("숙제 주제", placeholder="예: 과일 이름, 날씨 표현, 자기소개")
        hw_count = st.number_input("문제 수", min_value=5, max_value=30, value=10)

    hw_types = st.multiselect(
        "숙제 유형 선택 (복수 선택 가능)",
        ["📝 단어 받아쓰기", "❓ 빈칸 채우기", "🔄 한영 번역"],
        default=["📝 단어 받아쓰기", "❓ 빈칸 채우기", "🔄 한영 번역"]
    )

    if st.button("📄 숙제 생성하기", type="primary"):
        if not api_key:
            st.warning("왼쪽 사이드바에 API 키를 입력해주세요!")
        elif not hw_topic:
            st.warning("숙제 주제를 입력해주세요!")
        elif not hw_types:
            st.warning("숙제 유형을 하나 이상 선택해주세요!")
        else:
            with st.spinner("AI가 숙제를 만드는 중이에요... ✍️"):
                client = anthropic.Anthropic(api_key=api_key)
                prompt = f"""
당신은 초등학교 영어 전문 강사입니다.
아래 정보를 바탕으로 영어 숙제를 한국어와 영어로 작성해주세요.
- 대상: 초등학교 {hw_grade}
- 학생 수준: {hw_level}
- 숙제 주제: {hw_topic}
- 총 문제 수: {hw_count}개
- 숙제 유형: {', '.join(hw_types)}

다음 형식으로 작성해주세요:

[영어 숙제 학습지]
주제: {hw_topic}
학년: {hw_grade} / 수준: {hw_level}

선택된 유형별로 문제를 균등하게 나눠서 만들어주세요.

📝 단어 받아쓰기 문제가 포함된 경우:
- 한국어 뜻을 보고 영어 단어를 쓰는 형식
- 예) 사과 → ________

❓ 빈칸 채우기 문제가 포함된 경우:
- 영어 문장에서 단어가 빠진 형식
- 예) I like ________ . (나는 사과를 좋아해요)

🔄 한영 번역 문제가 포함된 경우:
- 한국어 문장을 영어로 쓰는 형식
- 예) 나는 학생이에요 → ________________

마지막에 정답지도 함께 작성해주세요!
"""
                message = client.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=3000,
                    messages=[{"role": "user", "content": prompt}]
                )
                result = message.content[0].text
                st.success("✅ 숙제가 완성됐어요!")
                st.divider()
                st.markdown(result)

                # 워드 파일 생성
                doc = Document()
                doc.add_heading(f"영어 숙제 - {hw_topic}", 0)
                doc.add_paragraph(f"학년: {hw_grade}  |  수준: {hw_level}")
                doc.add_paragraph("")
                doc.add_paragraph(result)
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="📥 숙제 워드파일 다운로드 (.docx)",
                    data=buffer,
                    file_name=f"영어숙제_{hw_topic}_{hw_grade}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )